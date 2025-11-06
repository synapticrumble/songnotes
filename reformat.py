import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


def set_georgia_font(run):
    """Set Georgia font for optimal HTML rendering"""
    run.font.name = 'Georgia'
    run.font.size = None  # Use default size


def add_page_break(paragraph):
    """Add a page break after the paragraph"""
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_hyperlink(paragraph, bookmark_name, display_text, is_bold=False, is_underlined=True):
    """Add a clickable hyperlink to a bookmark"""
    run = paragraph.add_run()
    
    # Create hyperlink field
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.text = f'HYPERLINK \\l "{bookmark_name}"'
    
    fld_char_sep = OxmlElement('w:fldChar')
    fld_char_sep.set(qn('w:fldCharType'), 'separate')
    
    # Add the display text
    run_text = paragraph.add_run(display_text)
    set_georgia_font(run_text)
    if is_bold:
        run_text.bold = True
    if is_underlined:
        run_text.underline = True
        run_text.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
    
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    
    # Insert field elements
    r_element = run._r
    r_element.append(fld_char_begin)
    r_element.append(instr_text)
    r_element.append(fld_char_sep)
    
    # End field after text
    end_run = paragraph.add_run()
    end_run._r.append(fld_char_end)


def add_bookmark(paragraph, bookmark_name, bookmark_id):
    """Add a named bookmark to a paragraph"""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    paragraph._p.insert(0, start)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.append(end)


def insert_paragraph_after(paragraph, text=""):
    """Insert a new paragraph after the given paragraph"""
    # Get the paragraph's parent (document)
    parent = paragraph._parent
    # Create a new paragraph
    new_para = parent.add_paragraph(text)
    # Move it to the correct position by manipulating the XML
    parent._element.remove(new_para._element)
    paragraph._element.addnext(new_para._element)
    return new_para


def is_separator(text):
    """Check if text is a separator line (=====****=====)"""
    return bool(re.fullmatch(r'=+\*+=+', text.strip()))


def process_docx(input_file, output_file):
    """Process the document to add TOC, bookmarks, and links"""
    doc = Document(input_file)
    
    # STEP 1: Store paragraph OBJECTS before modifying document structure
    all_paragraphs = list(doc.paragraphs)  # Create a copy of paragraph objects
    song_data = []  # Store (separator_para, title_para, title_text) tuples
    
    # Collect song information while preserving paragraph objects
    # Enhanced to find the first non-empty line as song title after separators
    for i, para in enumerate(all_paragraphs):
        if is_separator(para.text):
            # Look for the first non-empty line after the separator
            title_para = None
            title_text = ""
            
            # Search through subsequent paragraphs to find the first non-empty title
            for j in range(i + 1, len(all_paragraphs)):
                candidate_para = all_paragraphs[j]
                candidate_text = candidate_para.text.strip()
                
                # If we find non-empty text, this is our song title
                if candidate_text:
                    title_para = candidate_para
                    title_text = candidate_text
                    break
                
                # Stop searching after too many empty lines (max 5 empty lines)
                if j - i > 5:
                    break
            
            # Only add to song_data if we found a valid title
            if title_para and title_text:
                song_data.append((para, title_para, title_text))
    
    # STEP 2: Create Table of Contents at the beginning
    # Insert TOC before the first paragraph
    first_para = doc.paragraphs[0]
    
    # Create TOC header
    toc_header = first_para.insert_paragraph_before("Table of Contents")
    toc_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_header_run = toc_header.runs[0]
    toc_header_run.bold = True
    set_georgia_font(toc_header_run)
    
    # Add bookmark for "Top"
    add_bookmark(toc_header, "Top", 0)
    
    # Add empty line after TOC header
    empty_line = insert_paragraph_after(toc_header, "")
    
    # STEP 3: Add TOC entries with clickable links (chronological order)
    current_para = empty_line
    for idx, (_, _, title_text) in enumerate(song_data, 1):
        toc_entry = insert_paragraph_after(current_para, "")
        toc_entry.alignment = WD_ALIGN_PARAGRAPH.LEFT
        bookmark_name = f"song_{idx}"
        add_hyperlink(toc_entry, bookmark_name, f"{idx}. {title_text}", is_bold=True)
        current_para = toc_entry
    
    # Add empty line after TOC
    insert_paragraph_after(current_para, "")
    
    # STEP 4: Process songs - add bookmarks, page breaks, and back-to-top links
    bookmark_id = 1
    for idx, (separator_para, title_para, title_text) in enumerate(song_data, 1):
        # Add page break after separator line
        add_page_break(separator_para)
        
        # Add bookmark to title paragraph
        bookmark_name = f"song_{idx}"
        add_bookmark(title_para, bookmark_name, bookmark_id)
        
        # Set Georgia font for title
        for run in title_para.runs:
            set_georgia_font(run)
        if not title_para.runs:  # If no runs exist, create one
            run = title_para.add_run(title_para.text)
            title_para.clear()
            title_para.add_run(title_para.text)
            set_georgia_font(title_para.runs[0])
        
        # Add "Back to Top" link after title
        back_to_top_para = insert_paragraph_after(title_para, "")
        add_hyperlink(back_to_top_para, "Top", "Back to Top", is_underlined=True)
        
        bookmark_id += 1
    
    # STEP 5: Apply Georgia font to all text in the document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_georgia_font(run)
    
    # Save the processed document
    doc.save(output_file)
    print(f"Processed document saved as: {output_file}")
    print(f"Added {len(song_data)} songs to Table of Contents")



# Usage
if __name__ == "__main__":
    input_file = r"P:\\ShareDownloads\\BansuriMusic.docx"
    output_file = r"P:\\ShareDownloads\\songnotes\\songs_reformatted.docx"
    process_docx(input_file, output_file)
    print("File converted successfully!")